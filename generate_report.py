"""
generate_report.py - Generates a Word (.docx) progress report for the FYP project.
Includes objectives, methodology, graph analysis, findings, and next steps.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import DATA_DIR, OUTPUT_DIR, PLOTS_DIR
from data_loader import load_all_data, pair_on_off


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading


def add_paragraph_text(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        # Clear default and add fresh
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_image_with_caption(doc, image_path, caption, width=Inches(5.5)):
    if not os.path.exists(image_path):
        add_paragraph_text(doc, f"[Image not found: {image_path}]", italic=True)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    cap.paragraph_format.space_after = Pt(12)


def add_table_from_data(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def generate_word_report(data_dir, plots_dir, output_dir):
    # Load data
    df = load_all_data(data_dir)
    pairs = pair_on_off(df)

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =============================================
    # TITLE PAGE
    # =============================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Powered Parameter Optimization for\nGraphene Metamaterial THz Devices")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("FYP Progress Report - Phase 1: Data Analysis")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    info_lines = [
        "Student: Linah Salem Alrejaee",
        "Student ID: 231001766",
        "Programme: BSc Computer Science",
        "Supervisor: Dr. Riccardo Degl'Innocenti",
        "Institution: Queen Mary University of London",
        "Date: March 2026",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    doc.add_page_break()

    # =============================================
    # TABLE OF CONTENTS (manual)
    # =============================================
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Project Objectives",
        "2. Methodology - What Was Done in This Phase",
        "3. Dataset Overview",
        "4. Graph Analysis and Findings",
        "    4.1 All S12 Transmission Curves",
        "    4.2 Effect of Capacitor Width (c_w) on S12 Dip",
        "    4.3 Effect of Unit Cell Size (dx) on S12 Dip",
        "    4.4 ON vs OFF State Comparison",
        "    4.5 Frequency Shift and Dip Depth Summary",
        "5. Key Findings and Interpretation",
        "6. Current Status - What Has Been Done So Far",
        "7. Next Steps - What Will Be Done",
        "8. Complete Project Plan",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # =============================================
    # 1. PROJECT OBJECTIVES
    # =============================================
    add_heading_styled(doc, "1. Project Objectives", level=1)

    add_paragraph_text(doc,
        "This project aims to develop a Machine Learning tool that predicts and optimises the "
        "performance of graphene-based terahertz (THz) metamaterial modulators. These devices are "
        "simulated in COMSOL Multiphysics, where each simulation takes approximately 24 hours. "
        "The goal is to replace this brute-force approach with an ML model that can predict device "
        "performance instantly.")

    add_heading_styled(doc, "Primary Objectives", level=2)

    add_bullet(doc, "Maximise S12/S21 dip depth: ", bold_prefix="Objective 1: ")
    add_paragraph_text(doc,
        "The S12 parameter measures signal transmission through the device. A deeper dip (more "
        "negative dB value) indicates stronger resonance. Current best is -13.41 dB; the target "
        "is -20 dB or deeper.", font_size=11)

    add_bullet(doc, "Maximise frequency shift between ON and OFF states: ", bold_prefix="Objective 2: ")
    add_paragraph_text(doc,
        "The device operates in two states controlled by graphene conductivity "
        "(sigma=0.3 mS for ON, sigma=1.2 mS for OFF). The frequency at which the S12 dip occurs "
        "shifts between these states. A larger shift means better tunability. Current best is 60 GHz.",
        font_size=11)

    add_heading_styled(doc, "The 7 Tuneable Parameters", level=2)

    param_headers = ["Parameter", "Description", "Current Range", "Unit"]
    param_rows = [
        ["g_w", "Capacitor gap width", "3 - 5", "um"],
        ["c_w", "Capacitor width", "12 - 28", "um"],
        ["dx", "Unit cell size", "23 - 35", "um"],
        ["w_au", "Gold line width", "4 (fixed)", "um"],
        ["w_au_top", "Top gate bar thickness", "1.5 (fixed)", "um"],
        ["w_graph", "Graphene extra width", "1 (fixed)", "um"],
        ["h_graph", "Graphene extra height", "2 (fixed)", "um"],
    ]
    add_table_from_data(doc, param_headers, param_rows)

    doc.add_page_break()

    # =============================================
    # 2. METHODOLOGY
    # =============================================
    add_heading_styled(doc, "2. Methodology - What Was Done in This Phase", level=1)

    add_paragraph_text(doc,
        "In this initial phase, the focus was on building the data processing pipeline and "
        "performing exploratory data analysis on the available COMSOL simulation results.")

    add_heading_styled(doc, "Steps Completed", level=2)

    steps = [
        ("Data Collection: ", "Received 13 unique COMSOL simulation data files from supervisor, "
         "each containing S12 and S22 parameter curves across a frequency range of 300-800 GHz."),
        ("Data Cleaning: ", "Identified and removed 12 duplicate files (Windows download copies). "
         "Cleaned dataset contains 13 unique simulations."),
        ("Pipeline Development: ", "Built a Python data processing pipeline (data_loader.py) that "
         "automatically parses COMSOL export files, extracts simulation parameters from filenames, "
         "identifies S12 dips, and pairs ON/OFF states."),
        ("Visualisation: ", "Created 5 analysis plots (visualize.py) examining parameter effects "
         "on S12 dip depth and frequency shift."),
        ("Analysis: ", "Generated summary statistics and parameter sensitivity analysis (analyze.py) "
         "to understand which parameters have the strongest influence on device performance."),
    ]
    for prefix, text in steps:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_styled(doc, "Tools Used", level=2)
    tools = [
        "Python 3.11",
        "pandas - data manipulation and analysis",
        "NumPy - numerical computations",
        "Matplotlib - graph generation",
        "python-docx - automated report generation",
    ]
    for t in tools:
        add_bullet(doc, t)

    doc.add_page_break()

    # =============================================
    # 3. DATASET OVERVIEW
    # =============================================
    add_heading_styled(doc, "3. Dataset Overview", level=1)

    add_paragraph_text(doc,
        f"The dataset consists of {len(df)} unique COMSOL simulations. Each file contains "
        f"S-parameter data (S12 transmission and S22 reflection) across 26 frequency points "
        f"from 300 GHz to 800 GHz. The simulations vary three geometric parameters (dx, g_w, c_w) "
        f"and the graphene conductivity (sigma).")

    add_heading_styled(doc, "Complete Simulation Table", level=2)

    sim_headers = ["#", "dx (um)", "g_w (um)", "c_w (um)", "sigma (mS)", "S12 Dip (dB)", "Dip Freq (GHz)"]
    sim_rows = []
    for i, (_, row) in enumerate(df.iterrows()):
        sim_rows.append([
            str(i + 1),
            str(int(row['dx'])),
            str(int(row['g_w'])),
            str(int(row['c_w'])),
            str(row['sigma']),
            f"{row['s12_min']:.2f}",
            f"{row['s12_min_freq']:.0f}",
        ])
    add_table_from_data(doc, sim_headers, sim_rows)

    add_heading_styled(doc, "ON/OFF Pair Summary", level=2)

    add_paragraph_text(doc,
        f"Out of 13 simulations, {len(pairs)} complete ON/OFF pairs were identified "
        f"(matching geometric parameters with both sigma=0.3 and sigma=1.2).")

    if len(pairs) > 0:
        pair_headers = ["dx", "g_w", "c_w", "ON Dip (dB)", "ON Freq (GHz)",
                        "OFF Dip (dB)", "OFF Freq (GHz)", "Freq Shift (GHz)"]
        pair_rows = []
        for _, p in pairs.iterrows():
            pair_rows.append([
                str(int(p['dx'])), str(int(p['g_w'])), str(int(p['c_w'])),
                f"{p['s12_min_on']:.2f}", f"{p['s12_min_freq_on']:.0f}",
                f"{p['s12_min_off']:.2f}", f"{p['s12_min_freq_off']:.0f}",
                f"{p['freq_shift_ghz']:.0f}",
            ])
        add_table_from_data(doc, pair_headers, pair_rows)

    doc.add_page_break()

    # =============================================
    # 4. GRAPH ANALYSIS
    # =============================================
    add_heading_styled(doc, "4. Graph Analysis and Findings", level=1)

    # --- 4.1 All S12 curves ---
    add_heading_styled(doc, "4.1 All S12 Transmission Curves", level=2)
    add_image_with_caption(doc,
        os.path.join(plots_dir, '1_all_s12_curves.png'),
        "Figure 1: All S12 transmission curves. Blue = ON state (sigma=0.3), Red = OFF state (sigma=1.2)")

    add_paragraph_text(doc,
        "Figure 1 shows all 13 S12 transmission curves overlaid. The blue curves represent the "
        "ON state (graphene conductivity sigma=0.3 mS) and the red curves represent the OFF state "
        "(sigma=1.2 mS). Key observations:")

    observations_1 = [
        "Each curve shows a clear resonance dip where transmission is minimised, indicating "
        "strong electromagnetic absorption by the metamaterial structure.",
        "The dip frequency shifts depending on the geometric parameters - larger capacitor width "
        "(c_w) and larger unit cell size (dx) push the resonance to lower frequencies.",
        "The ON and OFF states show visibly different dip positions, confirming that the graphene "
        "conductivity effectively tunes the resonance frequency.",
        "All dips fall in the range of -9.6 to -13.4 dB, still significantly above the -20 dB target.",
    ]
    for obs in observations_1:
        add_bullet(doc, obs)

    # --- 4.2 S12 vs c_w ---
    add_heading_styled(doc, "4.2 Effect of Capacitor Width (c_w) on S12 Dip Depth", level=2)
    add_image_with_caption(doc,
        os.path.join(plots_dir, '2_s12_dip_vs_cw.png'),
        "Figure 2: S12 dip depth vs capacitor width (c_w) for g_w=3 and g_w=5 at dx=35")

    add_paragraph_text(doc,
        "Figure 2 isolates the effect of capacitor width (c_w) on the S12 dip depth, comparing "
        "two gap width values (g_w=3 and g_w=5). Analysis:")

    observations_2 = [
        "Increasing c_w from 12 to 28 um deepens the S12 dip for g_w=5 (from -12.0 to -12.6 dB), "
        "showing a clear positive correlation between capacitor width and resonance strength.",
        "The g_w=5 configuration (orange) consistently shows deeper dips than g_w=3 (blue) at c_w=28, "
        "suggesting that a wider capacitor gap improves resonance coupling.",
        "The sensitivity is approximately 0.040 dB per um of c_w increase - a moderate effect.",
        "Extrapolating this trend, even c_w=50 um would only reach approximately -13.5 dB, "
        "suggesting that c_w alone cannot achieve the -20 dB target. Other parameters must be explored.",
    ]
    for obs in observations_2:
        add_bullet(doc, obs)

    # --- 4.3 S12 vs dx ---
    add_heading_styled(doc, "4.3 Effect of Unit Cell Size (dx) on S12 Dip Depth", level=2)
    add_image_with_caption(doc,
        os.path.join(plots_dir, '3_s12_dip_vs_dx.png'),
        "Figure 3: S12 dip depth vs unit cell size (dx) for g_w=3, c_w=28")

    add_paragraph_text(doc,
        "Figure 3 shows the dx sweep data - six simulations where only the unit cell size varies "
        "from 23 to 35 um. This reveals a strong linear relationship:")

    observations_3 = [
        "Increasing dx from 23 to 35 um deepens the S12 dip from -11.1 to -12.5 dB, "
        "a change of 1.4 dB over 12 um.",
        "The sensitivity is 0.117 dB per um - approximately 3x stronger than the c_w effect. "
        "This makes dx the most influential parameter identified so far.",
        "The resonance frequency also shifts significantly: from 540 GHz (dx=23) down to 400 GHz (dx=35), "
        "a shift of 140 GHz across the sweep range (11.7 GHz per um).",
        "The near-linear trend suggests that further increasing dx beyond 35 um could yield even deeper "
        "dips, though physical constraints may limit this.",
    ]
    for obs in observations_3:
        add_bullet(doc, obs)

    # --- 4.4 ON vs OFF ---
    add_heading_styled(doc, "4.4 ON vs OFF State Comparison", level=2)
    add_image_with_caption(doc,
        os.path.join(plots_dir, '4_on_off_comparison.png'),
        "Figure 4: ON (blue) vs OFF (red) state S12 curves for all 4 complete pairs",
        width=Inches(6.0))

    add_paragraph_text(doc,
        "Figure 4 shows side-by-side comparisons of the 4 complete ON/OFF pairs. The green "
        "annotation shows the frequency shift between states:")

    observations_4 = [
        "Three configurations with g_w=5 show a consistent 60 GHz frequency shift between ON and OFF "
        "states, regardless of c_w value (12, 20, or 28 um). This suggests that the ON/OFF shift "
        "is primarily controlled by the graphene properties, not the capacitor geometry.",
        "The g_w=3, c_w=28 configuration shows a smaller shift of only 40 GHz, indicating that "
        "smaller gap width reduces the graphene's ability to tune the resonance.",
        "In the OFF state (sigma=1.2), the dips tend to be shallower (less negative) except for "
        "the g_w=3, c_w=28 case where the OFF dip (-13.4 dB) is actually deeper than the ON dip (-12.5 dB).",
        "The g_w=5, c_w=28 configuration offers the best trade-off: 60 GHz shift with the deepest "
        "average dip (-12.6 dB).",
    ]
    for obs in observations_4:
        add_bullet(doc, obs)

    # --- 4.5 Freq shift bar chart ---
    add_heading_styled(doc, "4.5 Frequency Shift and Dip Depth Summary", level=2)
    add_image_with_caption(doc,
        os.path.join(plots_dir, '5_freq_shift_and_dip.png'),
        "Figure 5: Frequency shift (bars) and average dip depth (diamonds) per configuration")

    add_paragraph_text(doc,
        "Figure 5 combines both optimisation objectives in a single chart. The green bars show "
        "the frequency shift, while the red diamonds show the average dip depth:")

    observations_5 = [
        "The three g_w=5 configurations all achieve 60 GHz frequency shift (the maximum observed), "
        "while the g_w=3 configuration only achieves 40 GHz.",
        "Among the 60 GHz shift configurations, g_w=5, c_w=28 has the deepest average dip (-12.6 dB), "
        "making it the best overall configuration found so far.",
        "There is a clear trade-off visible: g_w=5, c_w=12 has 60 GHz shift but the shallowest dip "
        "(-10.8 dB), while g_w=3, c_w=28 has the deepest dip (-12.9 dB) but smallest shift (40 GHz).",
        "The ideal configuration would combine deep dip AND large shift - this is the multi-objective "
        "optimisation challenge that the ML model will address.",
    ]
    for obs in observations_5:
        add_bullet(doc, obs)

    doc.add_page_break()

    # =============================================
    # 5. KEY FINDINGS
    # =============================================
    add_heading_styled(doc, "5. Key Findings and Interpretation", level=1)

    add_heading_styled(doc, "Parameter Sensitivity Ranking", level=2)

    sens_headers = ["Parameter", "Dip Sensitivity (dB/um)", "Freq Sensitivity (GHz/um)", "Impact"]
    sens_rows = [
        ["dx (unit cell size)", "0.117", "11.7", "Strongest effect"],
        ["c_w (capacitor width)", "0.040", "8.8", "Moderate effect"],
        ["g_w (gap width)", "Affects shift more than dip", "-", "Controls ON/OFF shift"],
        ["w_au, w_au_top, w_graph, h_graph", "Unknown - not yet varied", "-", "Need data"],
    ]
    add_table_from_data(doc, sens_headers, sens_rows)

    add_heading_styled(doc, "Key Conclusions from Phase 1", level=2)

    conclusions = [
        ("dx is the most influential parameter: ",
         "It has 3x the sensitivity of c_w on dip depth and also strongly controls resonance frequency. "
         "Larger dx values consistently give deeper dips."),
        ("g_w controls the ON/OFF shift: ",
         "g_w=5 gives 60 GHz shift while g_w=3 gives only 40 GHz. Wider gap allows graphene to "
         "have more influence on the resonance."),
        ("Current performance is far from target: ",
         "Best dip is -13.4 dB vs target of -20 dB (6.6 dB gap). This gap likely cannot be closed "
         "by varying only dx, g_w, and c_w - the unexplored parameters (w_au, w_au_top, w_graph, "
         "h_graph) must be investigated."),
        ("Multi-objective trade-off exists: ",
         "Deeper dips and larger frequency shifts partially conflict. The ML optimisation model will "
         "need to find the Pareto-optimal combinations."),
        ("Data is insufficient for ML: ",
         "Only 3 of 7 parameters have been varied, and only 13 simulations exist. At minimum 20-30 "
         "simulations with all parameters varied are needed to train a useful ML model."),
    ]
    for prefix, text in conclusions:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # =============================================
    # 6. CURRENT STATUS
    # =============================================
    add_heading_styled(doc, "6. Current Status - What Has Been Done So Far", level=1)

    done_items = [
        ("Understanding the Problem: ",
         "Studied the COMSOL model, understood the 7 tuneable parameters, the S-parameter outputs, "
         "and the two optimisation objectives (deeper dip + larger frequency shift)."),
        ("Data Pipeline Built: ",
         "Created data_loader.py that automatically parses COMSOL export files, extracts parameters "
         "from filenames, reads S-parameter curves, finds dips, and pairs ON/OFF states. This pipeline "
         "scales automatically as new data files are added."),
        ("Visualisation Tool Built: ",
         "Created visualize.py that generates 5 analysis plots examining parameter effects from "
         "different perspectives. All plots are saved as high-resolution PNGs."),
        ("Analysis Complete: ",
         "Created analyze.py that computes summary statistics, parameter sensitivity, identifies "
         "best configurations, flags missing data, and generates recommendations."),
        ("Report Generation Automated: ",
         "Created generate_report.py that produces this Word document automatically, embedding "
         "all graphs and analysis. Can be regenerated whenever new data is added."),
        ("Initial Findings: ",
         "Identified dx as the most influential parameter, confirmed that g_w=5 gives better "
         "ON/OFF shift, and established that current data is insufficient for ML training."),
    ]

    for prefix, text in done_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # =============================================
    # 7. NEXT STEPS
    # =============================================
    add_heading_styled(doc, "7. Next Steps - What Will Be Done", level=1)

    add_heading_styled(doc, "Immediate Next Steps (Week 1-2)", level=2)
    immediate = [
        "Request sigma=1.2 simulation data for the 5 missing dx sweep configurations from supervisor.",
        "Request simulations where w_au, w_au_top, w_graph, and h_graph are varied "
        "(these 4 parameters have not been explored at all yet).",
        "Explore larger dx values (>35 um) and larger c_w values (>28 um) to push dip depth further.",
    ]
    for item in immediate:
        add_bullet(doc, item)

    add_heading_styled(doc, "Phase 2: ML Model Development (Week 3-5)", level=2)
    phase2 = [
        "Once 20-30 simulations are available, build Gaussian Process Regression model "
        "(ideal for small datasets with uncertainty quantification).",
        "Build Random Forest and XGBoost baseline models for comparison.",
        "Implement cross-validation to ensure models generalise well despite small dataset.",
        "Evaluate models using R-squared, RMSE, and MAE metrics.",
    ]
    for item in phase2:
        add_bullet(doc, item)

    add_heading_styled(doc, "Phase 3: Optimisation (Week 5-7)", level=2)
    phase3 = [
        "Implement Bayesian Optimisation to intelligently suggest promising parameter combinations.",
        "Build multi-objective optimisation (NSGA-II) to find Pareto-optimal trade-offs between "
        "dip depth and frequency shift.",
        "Generate ranked list of top 10 recommended parameter combinations.",
    ]
    for item in phase3:
        add_bullet(doc, item)

    add_heading_styled(doc, "Phase 4: Validation and Tool (Week 7-10)", level=2)
    phase4 = [
        "Professor validates top 3-5 ML predictions by running actual COMSOL simulations.",
        "Build Streamlit web interface for interactive parameter exploration.",
        "Finalise report and prepare for presentation.",
    ]
    for item in phase4:
        add_bullet(doc, item)

    doc.add_page_break()

    # =============================================
    # 8. COMPLETE PROJECT PLAN
    # =============================================
    add_heading_styled(doc, "8. Complete Project Plan", level=1)

    plan_headers = ["Phase", "Task", "Status", "Timeline"]
    plan_rows = [
        ["Phase 0", "Understand problem and COMSOL model", "Completed", "Week 1"],
        ["Phase 0", "Receive simulation data from supervisor", "Completed (partial)", "Week 1-2"],
        ["Phase 1", "Data cleaning and duplicate removal", "Completed", "Week 2"],
        ["Phase 1", "Build data processing pipeline", "Completed", "Week 2"],
        ["Phase 1", "Exploratory data analysis and visualisation", "Completed", "Week 2"],
        ["Phase 1", "Parameter sensitivity analysis", "Completed", "Week 2"],
        ["Phase 1", "Generate progress report", "Completed", "Week 2"],
        ["Phase 1", "Request additional simulation data", "Next", "Week 3"],
        ["Phase 2", "Collect 20-30 simulation results", "Pending", "Week 3-4"],
        ["Phase 2", "Build ML models (GP, RF, XGBoost)", "Pending", "Week 4-5"],
        ["Phase 2", "Model evaluation and comparison", "Pending", "Week 5-6"],
        ["Phase 3", "Bayesian Optimisation implementation", "Pending", "Week 6-7"],
        ["Phase 3", "Multi-objective optimisation (Pareto)", "Pending", "Week 7"],
        ["Phase 3", "Generate parameter recommendations", "Pending", "Week 7-8"],
        ["Phase 4", "Supervisor validates ML predictions", "Pending", "Week 8-9"],
        ["Phase 4", "Build Streamlit web interface", "Pending", "Week 9-10"],
        ["Phase 4", "Final report and presentation", "Pending", "Week 10-12"],
    ]
    add_table_from_data(doc, plan_headers, plan_rows)

    # =============================================
    # SAVE
    # =============================================
    output_path = os.path.join(output_dir, 'FYP_Progress_Report_Phase1.docx')
    doc.save(output_path)
    print(f"\nWord report saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_word_report(DATA_DIR, PLOTS_DIR, OUTPUT_DIR)
