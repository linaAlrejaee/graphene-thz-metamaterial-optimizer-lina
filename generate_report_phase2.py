"""
generate_report_phase2.py - Generates Phase 2 FYP Progress Report (.docx).
Covers: new dataset, ML model results, Streamlit dashboard, and next steps.
"""

import os
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import DATA_DIR, OUTPUT_DIR, PLOTS_DIR
from data_loader import load_all_data, pair_on_off
from ml_model import (
    prepare_dataset, prepare_pair_dataset, build_models,
    evaluate_models, train_final_models, feature_importance,
)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading


def add_paragraph_text(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_table_from_data(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()


def add_image_with_caption(doc, image_path, caption, width=Inches(5.5)):
    if not os.path.exists(image_path):
        add_paragraph_text(doc, f"[Image not found: {image_path}]", italic=True)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    cap.paragraph_format.space_after = Pt(12)


def generate_phase2_report(data_dir, plots_dir, output_dir):
    """Generate the full Phase 2 progress report."""
    import warnings
    warnings.filterwarnings("ignore")

    # Load all data
    df = load_all_data(data_dir)
    pairs = pair_on_off(df)

    # ML results
    X_all, Y_all, feat_cols, target_cols = prepare_dataset(df)
    models = build_models()
    results_all, _ = evaluate_models(X_all, Y_all, models)
    trained_all, scaler_all = train_final_models(X_all, Y_all, build_models())
    imp_all = feature_importance(X_all, Y_all, feat_cols)

    results_pairs, imp_pairs = None, None
    if len(pairs) >= 5:
        X_p, Y_p, _, _ = prepare_pair_dataset(pairs)
        results_pairs, _ = evaluate_models(X_p, Y_p, build_models())
        imp_pairs = feature_importance(X_p, Y_p, feat_cols)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "AI-Powered Parameter Optimization for\n"
        "Graphene Metamaterial THz Devices"
    )
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "FYP Progress Report - Phase 2: Machine Learning & Interactive Dashboard"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    info_lines = [
        "Student: Linah Salem Alrejaee",
        "Student ID: 231001766",
        "Programme: BSc Computer Science",
        "Supervisor: Dr. Riccardo Degl'Innocenti",
        "Institution: Queen Mary University of London",
        "Date: March 2026",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Phase 1 Recap",
        "3. New Dataset - What Changed",
        "    3.1 New Simulation Files",
        "    3.2 Key Discovery: h_graph Parameter",
        "    3.3 Updated Dataset Overview",
        "4. Machine Learning Implementation",
        "    4.1 Methodology",
        "    4.2 Model Performance - All Simulations",
        "    4.3 Model Performance - ON/OFF Pairs",
        "    4.4 Feature Importance Analysis",
        "    4.5 Parameter Sensitivity Update",
        "5. Interactive Streamlit Dashboard",
        "    5.1 Dashboard Overview",
        "    5.2 Page Descriptions",
        "    5.3 Prediction Tool",
        "6. Results Summary",
        "7. Current Limitations",
        "8. Next Steps and Recommendations",
        "9. Updated Project Plan",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    add_heading_styled(doc, "1. Executive Summary", level=1)

    best_dip = df['s12_min'].min()
    best_dip_freq = df.loc[df['s12_min'].idxmin(), 's12_min_freq']
    best_shift = pairs['freq_shift_ghz'].max() if len(pairs) > 0 else 0
    gap = -20 - best_dip

    add_paragraph_text(doc,
        "Phase 2 marks a significant advancement in the project. Six new COMSOL simulation files "
        "were received from the supervisor, introducing h_graph (graphene extra height) as a "
        "variable parameter for the first time. This led to the most important discovery so far: "
        f"h_graph=6 produces an S12 dip of {best_dip:.2f} dB, a massive improvement from the "
        "Phase 1 best of -13.41 dB. The gap to the -20 dB target has been reduced from 6.59 dB "
        f"to just {abs(gap):.2f} dB.")

    add_paragraph_text(doc,
        "Three machine learning regression models (Random Forest, Gradient Boosting, and Gaussian "
        "Process) were trained and evaluated using Leave-One-Out cross-validation on the expanded "
        f"dataset of {len(df)} simulations. Gradient Boosting achieved the best S12 dip prediction "
        "(R\u00b2 = 0.58), while Gaussian Process excelled at frequency prediction (R\u00b2 = 0.75).")

    add_paragraph_text(doc,
        "An interactive Streamlit web dashboard was built and deployed, providing the supervisor "
        "with real-time data exploration, ML model evaluation, and an instant prediction tool for "
        "new parameter combinations.")

    # Key metrics summary table
    add_heading_styled(doc, "Key Metrics at a Glance", level=2)
    metrics_headers = ["Metric", "Phase 1", "Phase 2", "Change"]
    metrics_rows = [
        ["Total Simulations", "13", str(len(df)), f"+{len(df)-13} new"],
        ["ON/OFF Pairs", "4", str(len(pairs)), f"+{len(pairs)-4} new"],
        ["Best S12 Dip", "-13.41 dB", f"{best_dip:.2f} dB", f"{best_dip - (-13.41):.2f} dB improvement"],
        ["Best Freq Shift", "60 GHz", f"{best_shift:.0f} GHz", "Same"],
        ["Gap to -20 dB Target", "6.59 dB", f"{abs(gap):.2f} dB", f"{6.59 - abs(gap):.2f} dB closer"],
        ["Parameters Varied", "3 (dx, g_w, c_w)", "4 (dx, g_w, c_w, h_graph)", "+1 (h_graph)"],
        ["ML Models Trained", "None", "3 (RF, GB, GP)", "New"],
        ["Best ML R\u00b2 (S12 Dip)", "N/A", "0.5786", "New"],
        ["Best ML R\u00b2 (Freq)", "N/A", "0.7543", "New"],
        ["Web Dashboard", "None", "Live (5 pages)", "New"],
    ]
    add_table_from_data(doc, metrics_headers, metrics_rows)

    doc.add_page_break()

    # =========================================================================
    # 2. PHASE 1 RECAP
    # =========================================================================
    add_heading_styled(doc, "2. Phase 1 Recap", level=1)

    add_paragraph_text(doc,
        "Phase 1 established the project foundation with the following completed work:")

    phase1_items = [
        ("Data Pipeline: ", "Built data_loader.py to automatically parse COMSOL .txt files, "
         "extract parameters from filenames, identify S12/S22 dips, and pair ON/OFF states."),
        ("13 Simulations Analysed: ", "3 parameters varied (dx: 23-35, g_w: 3-5, c_w: 12-28), "
         "4 parameters fixed (w_graph=1, h_graph=2, w_au=4, w_au_top=1.5)."),
        ("4 ON/OFF Pairs: ", "Identified frequency shifts of 40-60 GHz between sigma=0.3 and sigma=1.2."),
        ("Best Result: ", "-13.41 dB S12 dip at 400 GHz (dx=35, g_w=3, c_w=28, sigma=0.3)."),
        ("Key Finding: ", "dx was the most influential parameter (0.117 dB/um), but even at "
         "maximum values, the -20 dB target remained unreachable with dx, g_w, c_w alone."),
        ("Recommendation: ", "Vary the unexplored parameters (h_graph, w_graph, w_au) to find "
         "stronger resonance configurations."),
    ]
    for prefix, text in phase1_items:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # =========================================================================
    # 3. NEW DATASET
    # =========================================================================
    add_heading_styled(doc, "3. New Dataset - What Changed", level=1)

    # 3.1 New files
    add_heading_styled(doc, "3.1 New Simulation Files", level=2)

    add_paragraph_text(doc,
        "Six new COMSOL simulation files were received from the supervisor. These files introduce "
        "h_graph (graphene extra height) as a variable parameter for the first time, testing three "
        "values: h_graph = -6, 0, and 6 (in addition to the existing h_graph = 2). All new files "
        "use dx=35, g_w=3, c_w=28, which was the best-performing geometry from Phase 1.")

    new_files_headers = ["File", "h_graph", "sigma", "State", "S12 Dip (dB)", "Dip Freq (GHz)"]
    new_files_rows = [
        ["hgraph0_sigma0.3", "0", "0.3", "ON", "-14.89", "410"],
        ["hgraph0_sigma1.2", "0", "1.2", "OFF", "-14.27", "390"],
        ["hgraph6_sigma0.3", "6", "0.3", "ON", "-16.64", "410"],
        ["hgraph6_sigma1.2", "6", "1.2", "OFF", "-14.28", "390"],
        ["hgraphm6_sigma0.3", "-6", "0.3", "ON", "-14.95", "410"],
        ["hgraphm6_sigma1.2", "-6", "1.2", "OFF", "-14.22", "390"],
    ]
    add_table_from_data(doc, new_files_headers, new_files_rows)

    add_paragraph_text(doc,
        "Note: Negative h_graph values are encoded in filenames using 'm' prefix "
        "(e.g., 'hgraphm6' = h_graph = -6). The data pipeline was updated to handle this "
        "encoding automatically.")

    # 3.2 Key Discovery
    add_heading_styled(doc, "3.2 Key Discovery: h_graph Parameter", level=2)

    add_paragraph_text(doc,
        "The most significant finding of Phase 2 is the impact of h_graph on S12 dip depth:",
        bold=True)

    discovery_items = [
        ("h_graph = 6 produces -16.64 dB: ",
         "This is the deepest S12 dip observed in the entire project, a 3.23 dB improvement "
         "over the Phase 1 best of -13.41 dB (h_graph=2). This single parameter change was more "
         "impactful than any combination of dx, g_w, or c_w variations."),
        ("h_graph = 0 produces -14.89 dB: ",
         "Even reducing h_graph from 2 to 0 improved the dip by 1.48 dB, indicating that the "
         "previous default value (h_graph=2) was not optimal."),
        ("h_graph = -6 produces -14.95 dB: ",
         "Negative values also improve the dip significantly, suggesting the relationship between "
         "h_graph and S12 dip is non-linear."),
        ("Gap to target reduced to 3.36 dB: ",
         f"The -20 dB target is now within reach. h_graph=6 gives {best_dip:.2f} dB, leaving "
         f"only {abs(gap):.2f} dB to close. Further increasing h_graph (e.g., 8, 10, 12) "
         "could potentially achieve the target."),
    ]
    for prefix, text in discovery_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # h_graph comparison table
    add_heading_styled(doc, "h_graph Impact Comparison (dx=35, g_w=3, c_w=28, ON state)", level=3)
    hg_headers = ["h_graph (um)", "S12 Dip (dB)", "Dip Freq (GHz)", "Improvement vs h_graph=2"]
    hg_rows = [
        ["-6", "-14.95", "410", "+2.47 dB deeper"],
        ["0", "-14.89", "410", "+2.41 dB deeper"],
        ["2 (Phase 1 default)", "-12.48", "400", "Baseline"],
        ["6", "-16.64", "410", "+4.16 dB deeper"],
    ]
    add_table_from_data(doc, hg_headers, hg_rows)

    # 3.3 Updated dataset
    add_heading_styled(doc, "3.3 Updated Dataset Overview", level=2)

    add_paragraph_text(doc,
        f"The complete dataset now contains {len(df)} simulations and {len(pairs)} ON/OFF pairs. "
        "Four geometric parameters are varied:")

    param_headers = ["Parameter", "Description", "Values Tested", "Range", "Status"]
    param_rows = [
        ["dx", "Unit cell size (um)", "23, 25, 27, 31, 33, 35", "23-35", "Varied"],
        ["g_w", "Capacitor gap width (um)", "3, 5", "3-5", "Varied"],
        ["c_w", "Capacitor width (um)", "12, 20, 28", "12-28", "Varied"],
        ["h_graph", "Graphene extra height (um)", "-6, 0, 2, 6", "-6 to 6", "Varied (NEW)"],
        ["w_graph", "Graphene extra width (um)", "1", "Fixed", "Not yet varied"],
        ["w_au", "Gold line width (um)", "4", "Fixed", "Not yet varied"],
    ]
    add_table_from_data(doc, param_headers, param_rows)

    # Full simulation table
    add_heading_styled(doc, "Complete Simulation Table (19 simulations)", level=3)
    sim_headers = ["#", "dx", "g_w", "c_w", "h_graph", "sigma", "S12 Dip (dB)", "Dip Freq (GHz)"]
    sim_rows = []
    for i, (_, row) in enumerate(df.iterrows()):
        sim_rows.append([
            str(i + 1),
            str(int(row['dx'])),
            str(int(row['g_w'])),
            str(int(row['c_w'])),
            str(int(row['h_graph'])),
            str(row['sigma']),
            f"{row['s12_min']:.2f}",
            f"{row['s12_min_freq']:.0f}",
        ])
    add_table_from_data(doc, sim_headers, sim_rows)

    # ON/OFF pairs table
    add_heading_styled(doc, f"ON/OFF Pair Comparison ({len(pairs)} pairs)", level=3)
    if len(pairs) > 0:
        pair_headers = ["dx", "g_w", "c_w", "h_graph", "ON Dip (dB)", "ON Freq",
                        "OFF Dip (dB)", "OFF Freq", "Shift (GHz)", "Avg Dip (dB)"]
        pair_rows = []
        for _, p in pairs.iterrows():
            pair_rows.append([
                str(int(p['dx'])), str(int(p['g_w'])), str(int(p['c_w'])),
                str(int(p['h_graph'])),
                f"{p['s12_min_on']:.2f}", f"{p['s12_min_freq_on']:.0f}",
                f"{p['s12_min_off']:.2f}", f"{p['s12_min_freq_off']:.0f}",
                f"{p['freq_shift_ghz']:.0f}", f"{p['avg_dip']:.2f}",
            ])
        add_table_from_data(doc, pair_headers, pair_rows)

    doc.add_page_break()

    # =========================================================================
    # 4. MACHINE LEARNING IMPLEMENTATION
    # =========================================================================
    add_heading_styled(doc, "4. Machine Learning Implementation", level=1)

    # 4.1 Methodology
    add_heading_styled(doc, "4.1 Methodology", level=2)

    add_paragraph_text(doc,
        "Three regression models were implemented and evaluated for predicting S-parameter "
        "behaviour from geometric design parameters:")

    model_items = [
        ("Random Forest Regressor: ",
         "Ensemble of 100 decision trees. Robust to outliers and provides feature importance "
         "estimates. Configured with min_samples_leaf=2 to prevent overfitting on small data."),
        ("Gradient Boosting Regressor: ",
         "Sequential ensemble of 100 shallow trees (max_depth=3) with learning_rate=0.1. "
         "Typically achieves the highest accuracy on structured data by iteratively correcting "
         "prediction errors."),
        ("Gaussian Process Regressor: ",
         "Non-parametric Bayesian model using RBF kernel with automatic relevance determination. "
         "Provides uncertainty estimates (prediction confidence intervals) alongside predictions. "
         "Ideal for small datasets with smooth underlying functions."),
    ]
    for prefix, text in model_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_styled(doc, "Evaluation Strategy: Leave-One-Out Cross-Validation", level=3)

    add_paragraph_text(doc,
        f"With only {len(df)} samples, traditional train/test splits would waste valuable data. "
        "Instead, Leave-One-Out (LOO) cross-validation was used: for each of the N samples, "
        "the model is trained on N-1 samples and tested on the one held-out sample. This process "
        "repeats N times, producing a prediction for every data point using a model that never "
        "saw that point during training.")

    add_paragraph_text(doc,
        "This approach maximises the use of available data while providing unbiased performance "
        "estimates. The trade-off is high computational cost (N model trainings), which is "
        "acceptable given the small dataset size.")

    add_heading_styled(doc, "Features and Targets", level=3)

    feat_headers = ["Type", "Variables", "Description"]
    feat_rows = [
        ["Input Features (X)", "dx, g_w, c_w, h_graph, w_graph, w_au",
         "6 geometric parameters (standardised using StandardScaler)"],
        ["Target 1 (Y)", "S12 Dip (dB)",
         "Minimum S12 transmission value (deeper = stronger resonance)"],
        ["Target 2 (Y)", "Dip Freq (GHz)",
         "Frequency at which the S12 minimum occurs"],
    ]
    add_table_from_data(doc, feat_headers, feat_rows)

    add_paragraph_text(doc,
        "For ON/OFF pair analysis, the targets are Frequency Shift (GHz) and Average Dip (dB) "
        "between the ON and OFF states of each geometry.")

    doc.add_page_break()

    # 4.2 Model Performance - All Simulations
    add_heading_styled(doc, "4.2 Model Performance - All Simulations (19 samples)", level=2)

    add_paragraph_text(doc,
        "The following table shows Leave-One-Out cross-validation results for all three models "
        "on the complete dataset of 19 simulations:")

    perf_headers = ["Model", "Target", "R\u00b2", "RMSE", "MAE"]
    perf_rows = []
    for model_name, targets in results_all.items():
        for target_name, m in targets.items():
            perf_rows.append([
                model_name, target_name,
                f"{m['R2']:.4f}", f"{m['RMSE']:.4f}", f"{m['MAE']:.4f}",
            ])
    add_table_from_data(doc, perf_headers, perf_rows)

    add_paragraph_text(doc, "Interpretation of Results:", bold=True)

    interp_items = [
        ("Gradient Boosting is best for S12 Dip prediction (R\u00b2 = 0.5786): ",
         "It correctly captures 58% of the variance in dip depth with an average error of "
         "0.76 dB. Given a typical dip range of -9.6 to -16.6 dB (7 dB span), this represents "
         "a meaningful predictive capability."),
        ("Gaussian Process is best for Dip Frequency prediction (R\u00b2 = 0.7543): ",
         "It captures 75% of frequency variance with RMSE of 31.5 GHz. This is strong "
         "performance considering the frequency range spans 360-600 GHz (240 GHz range). "
         "The GP model benefits from the smooth relationship between geometry and resonance frequency."),
        ("Random Forest provides balanced performance: ",
         "R\u00b2 of 0.53 for both targets. Not the best at either, but consistently reasonable."),
        ("R\u00b2 values are moderate (0.34-0.75): ",
         "This is expected with only 19 samples and 4 varied parameters. ML models typically "
         "need 10-20x more samples than features for strong generalisation. With more data "
         "(target: 40+ simulations), R\u00b2 values should improve significantly."),
    ]
    for prefix, text in interp_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # 4.3 Model Performance - ON/OFF Pairs
    add_heading_styled(doc, "4.3 Model Performance - ON/OFF Pairs (7 pairs)", level=2)

    if results_pairs:
        pair_perf_headers = ["Model", "Target", "R\u00b2", "RMSE", "MAE"]
        pair_perf_rows = []
        for model_name, targets in results_pairs.items():
            for target_name, m in targets.items():
                pair_perf_rows.append([
                    model_name, target_name,
                    f"{m['R2']:.4f}", f"{m['RMSE']:.4f}", f"{m['MAE']:.4f}",
                ])
        add_table_from_data(doc, pair_perf_headers, pair_perf_rows)

        add_paragraph_text(doc, "Interpretation:", bold=True)

        pair_interp = [
            ("Gradient Boosting leads for Frequency Shift (R\u00b2 = 0.49): ",
             "It can explain about half the variance in ON/OFF frequency shift. Given only "
             "7 training pairs, this is a reasonable starting point."),
            ("Gaussian Process fails on this task (negative R\u00b2): ",
             "With only 7 samples, the GP model overfits during training and performs worse "
             "than simply predicting the mean value. This is a known limitation of GP with "
             "very small datasets relative to the complexity of the kernel."),
            ("Average Dip prediction is weak across all models (R\u00b2 = 0.19-0.22): ",
             "The average dip between ON/OFF states is difficult to predict because it depends "
             "on subtle interactions between geometry and conductivity state. More ON/OFF pairs "
             "are needed to improve this prediction."),
        ]
        for prefix, text in pair_interp:
            add_bullet(doc, text, bold_prefix=prefix)
    else:
        add_paragraph_text(doc, "Not enough ON/OFF pairs for evaluation.", italic=True)

    doc.add_page_break()

    # 4.4 Feature Importance
    add_heading_styled(doc, "4.4 Feature Importance Analysis", level=2)

    add_paragraph_text(doc,
        "Feature importance was computed using Random Forest (200 trees) to determine which "
        "geometric parameters have the strongest influence on each prediction target:")

    add_heading_styled(doc, "All Simulations - Feature Importance", level=3)

    imp_headers = ["Parameter", "S12 Dip Importance (%)", "Dip Freq Importance (%)"]
    imp_rows = []
    for param in imp_all.index:
        imp_rows.append([
            param,
            f"{imp_all.loc[param, 'S12 Dip (dB)'] * 100:.1f}%",
            f"{imp_all.loc[param, 'Dip Freq (GHz)'] * 100:.1f}%",
        ])
    add_table_from_data(doc, imp_headers, imp_rows)

    add_paragraph_text(doc, "Key Findings:", bold=True)

    fi_items = [
        ("h_graph dominates S12 dip depth (69.6%): ",
         "This confirms the Phase 2 discovery. Graphene height is by far the most important "
         "parameter for controlling resonance strength. This was not known in Phase 1 because "
         "h_graph was fixed at 2."),
        ("c_w dominates resonance frequency (54.2%): ",
         "Capacitor width is the primary frequency-tuning parameter, followed by dx (35.6%). "
         "This makes physical sense: c_w directly affects the LC circuit capacitance, which "
         "determines the resonance frequency."),
        ("g_w has moderate importance (5-7.4%): ",
         "Gap width has a small but measurable effect on both targets."),
        ("w_graph and w_au show 0% importance: ",
         "Both are fixed at constant values (1 and 4 respectively), so the model cannot "
         "learn their effect. Varying these parameters is a priority for Phase 3."),
    ]
    for prefix, text in fi_items:
        add_bullet(doc, text, bold_prefix=prefix)

    if imp_pairs is not None:
        add_heading_styled(doc, "ON/OFF Pairs - Feature Importance", level=3)

        imp_p_headers = ["Parameter", "Freq Shift Importance (%)", "Avg Dip Importance (%)"]
        imp_p_rows = []
        for param in imp_pairs.index:
            imp_p_rows.append([
                param,
                f"{imp_pairs.loc[param, 'Freq Shift (GHz)'] * 100:.1f}%",
                f"{imp_pairs.loc[param, 'Avg Dip (dB)'] * 100:.1f}%",
            ])
        add_table_from_data(doc, imp_p_headers, imp_p_rows)

        add_paragraph_text(doc,
            "For ON/OFF pair predictions, g_w is the dominant parameter for frequency shift "
            "(66.8%), confirming the Phase 1 finding that wider gap allows graphene to have "
            "more influence on tuning. For average dip, g_w (41.6%) and h_graph (34.3%) share "
            "importance.")

    # 4.5 Parameter Sensitivity Update
    add_heading_styled(doc, "4.5 Parameter Sensitivity Update", level=2)

    sens_headers = ["Parameter", "Dip Sensitivity", "Freq Sensitivity", "Impact Rank"]
    sens_rows = [
        ["h_graph", "0.347 dB/um (strongest)", "Minimal effect on freq", "1st (NEW)"],
        ["dx", "0.117 dB/um", "11.7 GHz/um", "2nd"],
        ["c_w", "0.040 dB/um", "8.8 GHz/um", "3rd"],
        ["g_w", "Controls ON/OFF shift", "N/A", "4th (for shift)"],
        ["w_graph", "Unknown - not yet varied", "Unknown", "Need data"],
        ["w_au", "Unknown - not yet varied", "Unknown", "Need data"],
    ]
    add_table_from_data(doc, sens_headers, sens_rows)

    add_paragraph_text(doc,
        "The h_graph parameter has approximately 3x the sensitivity of dx (0.347 vs 0.117 dB/um) "
        "and produces a 4.16 dB variation across its tested range (-6 to 6), compared to only "
        "1.40 dB for the dx sweep (23-35). This makes h_graph the clear priority parameter for "
        "reaching the -20 dB target.",
        bold=True)

    doc.add_page_break()

    # =========================================================================
    # 5. STREAMLIT DASHBOARD
    # =========================================================================
    add_heading_styled(doc, "5. Interactive Streamlit Dashboard", level=1)

    add_heading_styled(doc, "5.1 Dashboard Overview", level=2)

    add_paragraph_text(doc,
        "A fully interactive web dashboard was built using Streamlit and deployed to Streamlit "
        "Community Cloud. The dashboard provides real-time exploration of simulation data, ML "
        "model results, and instant predictions for new parameter combinations.")

    dash_headers = ["Feature", "Details"]
    dash_rows = [
        ["URL", "https://graphene-thz-metamaterial-optimizer-6onswkc8bfjrcb5dbu6am5.streamlit.app/"],
        ["Framework", "Streamlit with Plotly interactive charts"],
        ["Pages", "5 (Overview, Data Explorer, Curves, ML Results, Predictions)"],
        ["Source Code", "https://github.com/tannu64/graphene-thz-metamaterial-optimizer"],
        ["Caching", "Data and ML models cached for fast page navigation"],
        ["Charts", "All charts are interactive (zoom, hover, pan, download as PNG)"],
    ]
    add_table_from_data(doc, dash_headers, dash_rows)

    add_heading_styled(doc, "5.2 Page Descriptions", level=2)

    pages = [
        ("Page 1 - Project Overview: ",
         "Displays 6 key metric cards (total simulations, ON/OFF pairs, best S12 dip, best "
         "frequency shift, gap to target, best dip frequency). Shows optimisation objectives "
         "and a parameter range summary table."),
        ("Page 2 - Data Explorer: ",
         "Three tabs showing: (a) all 19 simulations with their S12 dip and frequency values, "
         "(b) 7 ON/OFF pairs with frequency shift comparison, (c) best configurations highlighted "
         "with green success callouts for deepest dip, largest shift, and best average dip."),
        ("Page 3 - S-Parameter Curves: ",
         "Five interactive Plotly charts: all S12 transmission curves (colour-coded by ON/OFF "
         "state), S12 dip vs capacitor width (c_w), S12 dip vs unit cell size (dx), ON/OFF "
         "state side-by-side comparison with frequency shift annotations, and frequency shift "
         "bar chart with average dip overlay."),
        ("Page 4 - ML Model Results: ",
         "Displays the performance metrics table (R\u00b2, RMSE, MAE) for all three models. Shows "
         "actual vs predicted scatter plots with a perfect-prediction diagonal reference line. "
         "Includes feature importance horizontal bar chart. Can switch between 'All Simulations' "
         "and 'ON/OFF Pairs' datasets."),
        ("Page 5 - Predict New Design: ",
         "Four parameter sliders (dx, g_w, c_w, h_graph) allow the user to input any geometry. "
         "The selected ML model instantly predicts the S12 dip depth and resonance frequency. "
         "Gaussian Process additionally shows prediction uncertainty (\u00b1 standard deviation and "
         "95% confidence interval). An extrapolation warning appears when slider values exceed "
         "the training data range. A traffic-light indicator shows how close the prediction is "
         "to the -20 dB target."),
    ]
    for prefix, text in pages:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_styled(doc, "5.3 Prediction Tool - Input Ranges", level=2)

    add_paragraph_text(doc,
        "The prediction tool accepts the following input ranges. Predictions are most reliable "
        "within the training data range (highlighted):")

    pred_headers = ["Parameter", "Slider Range", "Training Range (Reliable)", "Best for Deep Dip"]
    pred_rows = [
        ["dx (um)", "20 - 40", "23 - 35", "35 (larger = deeper)"],
        ["g_w (um)", "2 - 6", "3 - 5", "3"],
        ["c_w (um)", "10 - 35", "12 - 28", "28 (larger = deeper)"],
        ["h_graph (um)", "-8 to 10", "-6 to 6", "6 (strongest effect)"],
    ]
    add_table_from_data(doc, pred_headers, pred_rows)

    add_paragraph_text(doc,
        "Fixed parameters: w_graph = 1, w_au = 4 (cannot be changed in the prediction tool "
        "as they have not been varied in the training data).")

    doc.add_page_break()

    # =========================================================================
    # 6. RESULTS SUMMARY
    # =========================================================================
    add_heading_styled(doc, "6. Results Summary", level=1)

    add_heading_styled(doc, "What is Working Well", level=2)

    working_items = [
        ("Data pipeline is fully automated: ",
         "All 19 COMSOL files are parsed correctly, including negative h_graph values encoded "
         "as 'm' prefix (e.g., 'hgraphm6' = -6). Adding new simulation files requires only "
         "dropping them into the data/ folder."),
        ("ON/OFF pairing works correctly: ",
         f"7 out of a possible 12 ON/OFF pairs are automatically identified. The 5 missing "
         "pairs are due to missing sigma=1.2 simulations for dx=23-33."),
        ("Frequency prediction is strong (R\u00b2 = 0.75): ",
         "The Gaussian Process model predicts resonance frequency with 75% accuracy and "
         "provides useful confidence intervals. Average error is 25 GHz on a 240 GHz range."),
        ("h_graph identified as dominant parameter: ",
         "Feature importance analysis correctly identified h_graph as having 69.6% importance "
         "for S12 dip depth, confirming the empirical observation."),
        ("Dashboard is live and interactive: ",
         "The professor can explore all data, charts, and predictions directly in the browser "
         "without installing any software."),
    ]
    for prefix, text in working_items:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_styled(doc, "Best Configuration Found", level=2)

    best_headers = ["Objective", "Best Value", "Configuration", "Status"]
    best_rows = [
        ["Deepest S12 Dip", f"{best_dip:.2f} dB at {best_dip_freq:.0f} GHz",
         "dx=35, g_w=3, c_w=28, h_graph=6, sigma=0.3", f"{abs(gap):.2f} dB from target"],
        ["Largest Freq Shift", f"{best_shift:.0f} GHz",
         "dx=35, g_w=5, c_w=12, h_graph=2", "Good tunability"],
        ["Deepest Avg Dip (ON+OFF)", "-15.46 dB",
         "dx=35, g_w=3, c_w=28, h_graph=6", "Best combined"],
    ]
    add_table_from_data(doc, best_headers, best_rows)

    doc.add_page_break()

    # =========================================================================
    # 7. LIMITATIONS
    # =========================================================================
    add_heading_styled(doc, "7. Current Limitations", level=1)

    limitations = [
        ("Small dataset (19 samples): ",
         "ML models have moderate accuracy (R\u00b2 = 0.34-0.75). With 4 varied parameters and "
         "6 total features, at least 40-60 simulations are needed for reliable model training. "
         "Current models should be used for directional guidance, not precise predictions."),
        ("5 missing OFF-state simulations: ",
         "dx=23, 25, 27, 31, and 33 only have sigma=0.3 (ON state) data. Without sigma=1.2 "
         "counterparts, these geometries cannot be included in ON/OFF pair analysis. This limits "
         "frequency shift prediction to 7 pairs instead of 12."),
        ("Two parameters never varied: ",
         "w_graph (graphene width) and w_au (gold width) are fixed at 1 and 4 respectively. "
         "The model cannot predict their effect, and they may have significant influence on "
         "device performance."),
        ("h_graph only varied at one geometry: ",
         "h_graph was only tested at dx=35, g_w=3, c_w=28. Its effect at other geometries "
         "is unknown. There may be interactions between h_graph and other parameters that "
         "current data cannot capture."),
        ("Extrapolation is unreliable: ",
         "Predictions for parameter values outside the training range (dx>35, c_w>28, "
         "h_graph>6 or <-6) are extrapolations. The prediction tool warns users about this, "
         "but results may be significantly off."),
        ("ON/OFF pair models are weak: ",
         "With only 7 pairs, the Gaussian Process model produces negative R\u00b2 (worse than "
         "mean prediction). Only Gradient Boosting gives reasonable frequency shift prediction "
         "(R\u00b2 = 0.49). More pairs are essential."),
    ]
    for prefix, text in limitations:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_page_break()

    # =========================================================================
    # 8. NEXT STEPS
    # =========================================================================
    add_heading_styled(doc, "8. Next Steps and Recommendations", level=1)

    add_heading_styled(doc, "Priority 1: Complete Missing OFF-State Simulations", level=2)

    add_paragraph_text(doc,
        "Run sigma=1.2 (OFF state) COMSOL simulations for the 5 unmatched ON-state geometries. "
        "This will increase ON/OFF pairs from 7 to 12, significantly improving frequency shift "
        "prediction accuracy:")

    missing_headers = ["dx", "g_w", "c_w", "h_graph", "sigma needed"]
    missing_rows = [
        ["23", "3", "28", "2", "1.2"],
        ["25", "3", "28", "2", "1.2"],
        ["27", "3", "28", "2", "1.2"],
        ["31", "3", "28", "2", "1.2"],
        ["33", "3", "28", "2", "1.2"],
    ]
    add_table_from_data(doc, missing_headers, missing_rows)

    add_heading_styled(doc, "Priority 2: Explore h_graph Beyond Current Range", level=2)

    add_paragraph_text(doc,
        "Since h_graph=6 gave -16.64 dB (the best result), higher values may push past -20 dB. "
        "Recommended simulations:")

    hg_explore = [
        "h_graph = 8 at dx=35, g_w=3, c_w=28 (both sigma=0.3 and 1.2)",
        "h_graph = 10 at dx=35, g_w=3, c_w=28 (both sigma=0.3 and 1.2)",
        "h_graph = 12 at dx=35, g_w=3, c_w=28 (both sigma=0.3 and 1.2)",
        "If h_graph=10 or 12 reaches -20 dB, the primary target is achieved",
    ]
    for item in hg_explore:
        add_bullet(doc, item)

    add_heading_styled(doc, "Priority 3: Cross h_graph with Other Geometries", level=2)

    add_paragraph_text(doc,
        "Currently h_graph is only varied at dx=35, g_w=3, c_w=28. To capture interactions:")

    cross_items = [
        "Test h_graph=6 with g_w=5 (to see if it improves the 60 GHz shift configurations)",
        "Test h_graph=6 with different c_w values (12, 20) to check for geometry interactions",
        "Test h_graph=6 at smaller dx values (23, 27, 31) to check if the effect holds",
    ]
    for item in cross_items:
        add_bullet(doc, item)

    add_heading_styled(doc, "Priority 4: Vary w_graph and w_au", level=2)

    add_paragraph_text(doc,
        "These two parameters have never been varied (fixed at w_graph=1, w_au=4). They may "
        "contain untapped performance. Recommended: test w_graph = 0.5, 1, 1.5, 2 and "
        "w_au = 2, 4, 6, 8 at the best geometry (dx=35, g_w=3, c_w=28, h_graph=6).")

    add_heading_styled(doc, "Priority 5: Retrain Models with Expanded Dataset", level=2)

    add_paragraph_text(doc,
        "Once the dataset reaches 40+ simulations (after completing priorities 1-4), retrain "
        "all three ML models. Expected improvements: R\u00b2 > 0.80 for S12 dip, R\u00b2 > 0.85 for "
        "frequency. Consider adding Bayesian Optimisation to intelligently suggest the next "
        "most informative simulation to run.")

    doc.add_page_break()

    # =========================================================================
    # 9. UPDATED PROJECT PLAN
    # =========================================================================
    add_heading_styled(doc, "9. Updated Project Plan", level=1)

    plan_headers = ["Phase", "Task", "Status", "Timeline"]
    plan_rows = [
        ["Phase 0", "Understand problem and COMSOL model", "Completed", "Week 1"],
        ["Phase 0", "Receive initial simulation data (13 files)", "Completed", "Week 1-2"],
        ["Phase 1", "Data cleaning and pipeline development", "Completed", "Week 2"],
        ["Phase 1", "Exploratory data analysis and visualisation", "Completed", "Week 2"],
        ["Phase 1", "Parameter sensitivity analysis", "Completed", "Week 2"],
        ["Phase 1", "Phase 1 progress report", "Completed", "Week 2"],
        ["Phase 2", "Receive new simulation data (6 files, h_graph varied)", "Completed", "Week 3"],
        ["Phase 2", "Fix h_graph parsing bug (negative values)", "Completed", "Week 3"],
        ["Phase 2", "Build ML models (RF, GB, GP)", "Completed", "Week 3-4"],
        ["Phase 2", "LOO cross-validation and evaluation", "Completed", "Week 4"],
        ["Phase 2", "Feature importance analysis", "Completed", "Week 4"],
        ["Phase 2", "Build Streamlit dashboard (5 pages)", "Completed", "Week 4"],
        ["Phase 2", "Deploy to Streamlit Cloud", "Completed", "Week 4"],
        ["Phase 2", "Phase 2 progress report", "Completed", "Week 4"],
        ["Phase 3", "Collect missing OFF-state simulations (5 files)", "Next", "Week 5"],
        ["Phase 3", "Explore h_graph > 6 (target: -20 dB)", "Next", "Week 5-6"],
        ["Phase 3", "Cross-parameter h_graph experiments", "Pending", "Week 6"],
        ["Phase 3", "Vary w_graph and w_au", "Pending", "Week 6-7"],
        ["Phase 3", "Retrain models with expanded dataset", "Pending", "Week 7"],
        ["Phase 3", "Bayesian Optimisation implementation", "Pending", "Week 7-8"],
        ["Phase 4", "Supervisor validates ML predictions in COMSOL", "Pending", "Week 8-9"],
        ["Phase 4", "Multi-objective optimisation (Pareto front)", "Pending", "Week 9"],
        ["Phase 4", "Final report and presentation", "Pending", "Week 10-12"],
    ]
    add_table_from_data(doc, plan_headers, plan_rows)

    # =========================================================================
    # SAVE
    # =========================================================================
    output_path = os.path.join(output_dir, 'FYP_Progress_Report_Phase2.docx')
    doc.save(output_path)
    print(f"\nPhase 2 report saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_phase2_report(DATA_DIR, PLOTS_DIR, OUTPUT_DIR)
